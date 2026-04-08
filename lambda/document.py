import base64
import re
import urllib.request
import xml.etree.ElementTree as ET
import zipfile
from io import BytesIO
from typing import Optional

import markdown
from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


# ── Template repair ───────────────────────────────────────────────────────────

def _recover_xml(content: bytes) -> bytes:
    """Re-parse malformed XML parts with lxml recover=True to auto-repair."""
    from lxml import etree as _etree
    try:
        _etree.fromstring(content)
        return content
    except _etree.XMLSyntaxError:
        pass
    try:
        parser = _etree.XMLParser(recover=True, encoding="utf-8")
        root = _etree.fromstring(content, parser)
        if root is None:
            return content
        return _etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    except Exception:
        return content


def _fix_rels_fragment_targets(content: bytes) -> bytes:
    """Add TargetMode='External' to bookmark hyperlinks missing it in .rels files."""
    _NS = "http://schemas.openxmlformats.org/package/2006/relationships"
    ET.register_namespace("", _NS)
    try:
        root = ET.fromstring(content)
    except ET.ParseError:
        return content
    changed = False
    for rel in root:
        if rel.get("Target", "").startswith("#") and "TargetMode" not in rel.attrib:
            rel.set("TargetMode", "External")
            changed = True
    if not changed:
        return content
    return ET.tostring(root, encoding="UTF-8", xml_declaration=True)


def _normalize_docx_bytes(data: bytes) -> bytes:
    """
    Repair the MAIA template's three packaging defects:
      1. Entire OPC package nested inside a subdirectory → strip prefix.
      2. XML parts with cross-structure hyperlinks (header3.xml) → lxml recover.
      3. .rels bookmark hyperlinks missing TargetMode='External' → patch.
    """
    with zipfile.ZipFile(BytesIO(data)) as zin:
        names = zin.namelist()
        if "[Content_Types].xml" in names:
            prefix = ""
        else:
            ct_path = next((n for n in names if n.endswith("[Content_Types].xml")), None)
            if ct_path is None:
                raise ValueError("Invalid .docx: no [Content_Types].xml found.")
            prefix = ct_path[: -len("[Content_Types].xml")]

        out = BytesIO()
        with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                name = info.filename
                if not name.startswith(prefix):
                    continue
                new_name = name[len(prefix):]
                if not new_name:
                    continue
                blob = zin.read(name)
                if new_name.endswith(".rels"):
                    blob = _fix_rels_fragment_targets(blob)
                elif new_name.endswith(".xml"):
                    blob = _recover_xml(blob)
                zout.writestr(new_name, blob)
        return out.getvalue()


# ── Style helpers ─────────────────────────────────────────────────────────────

def _collect_styles(doc: Document) -> dict:
    return {s.name.lower(): s.name for s in doc.styles}


def _best_style(available: dict, *candidates: str) -> Optional[str]:
    for name in candidates:
        if name:
            found = available.get(name.lower())
            if found:
                return found
    return None


# ── Cover-page placeholder helpers ───────────────────────────────────────────

def _para_text(p_elem) -> str:
    """Concatenate all w:t text in a paragraph XML element."""
    return "".join((t.text or "") for t in p_elem.iter(qn("w:t")))


def _set_para_text(p_elem, new_text: str) -> None:
    """
    Coalesce all w:r runs in a paragraph to a single text value,
    preserving the formatting (rPr) of the first run.
    """
    runs = [c for c in p_elem if c.tag == qn("w:r")]
    for child in p_elem:
        if child.tag == qn("w:hyperlink"):
            runs.extend(c for c in child if c.tag == qn("w:r"))
    if not runs:
        return
    t_elems = runs[0].findall(qn("w:t"))
    if t_elems:
        t_elems[0].text = new_text
        t_elems[0].set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        for extra in t_elems[1:]:
            runs[0].remove(extra)
    else:
        new_t = OxmlElement("w:t")
        new_t.text = new_text
        new_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        runs[0].append(new_t)
    for run in runs[1:]:
        for t in run.findall(qn("w:t")):
            t.text = ""


def _replace_run_text(root_elem, old: str, new: str) -> None:
    """Substring-replace *old* with *new* inside any single w:t element."""
    for t in root_elem.iter(qn("w:t")):
        if t.text and old in t.text:
            t.text = t.text.replace(old, new)


def _replace_date_runs(p_elem, date_str: str) -> None:
    """
    Find every xx/xx/20xx date pattern in a paragraph's w:t elements
    (possibly split across 5 consecutive runs) and replace with *date_str*.
    Safe to call on mixed paragraphs (e.g. footer with page numbers).
    """
    t_list = [t for t in p_elem.iter(qn("w:t"))]
    replaced = set()
    for i, t in enumerate(t_list):
        if i in replaced or (t.text or "").strip().lower() != "xx":
            continue
        if i + 4 >= len(t_list):
            continue
        segment = "".join((te.text or "") for te in t_list[i: i + 5])
        if re.match(r"^xx/xx/20\d{2}$", segment, re.IGNORECASE):
            t.text = date_str
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            for te in t_list[i + 1: i + 5]:
                te.text = ""
                replaced.add(id(te))


# ── XML / paragraph helpers ───────────────────────────────────────────────────

def _add_horizontal_rule(paragraph) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    if not url:
        paragraph.add_run(text)
        return
    try:
        r_id = paragraph.part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rStyle = OxmlElement("w:rStyle")
        rStyle.set(qn("w:val"), "Hyperlink")
        rPr.append(rStyle)
        new_run.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        new_run.append(t)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
    except Exception:
        run = paragraph.add_run(text)
        run.underline = True


def _fetch_image(src: str) -> Optional[bytes]:
    if src.startswith("data:"):
        m = re.match(r"data:[^;]+;base64,(.+)", src, re.DOTALL)
        if m:
            return base64.b64decode(m.group(1))
    elif src.startswith(("http://", "https://")):
        try:
            req = urllib.request.Request(src, headers={"User-Agent": "Mozilla/5.0"})
            with urllib.request.urlopen(req, timeout=5) as resp:
                return resp.read()
        except Exception:
            return None
    return None


# ── Inline content ────────────────────────────────────────────────────────────

def _add_run(paragraph, text, bold=False, italic=False,
             underline=False, strike=False, code=False) -> None:
    if not text:
        return
    run = paragraph.add_run(text)
    if bold:      run.bold = True
    if italic:    run.italic = True
    if underline: run.underline = True
    if strike:    run.font.strike = True
    if code:
        run.font.name = "Courier New"
        run.font.size = Pt(9)


def _inline(paragraph, element, bold=False, italic=False,
            underline=False, strike=False, code=False) -> None:
    for child in element.children:
        if isinstance(child, NavigableString):
            _add_run(paragraph, str(child), bold, italic, underline, strike, code)
            continue
        if not hasattr(child, "name") or child.name is None:
            continue
        tag = child.name.lower()
        if tag in ("strong", "b"):
            _inline(paragraph, child, bold=True,    italic=italic,    underline=underline, strike=strike, code=code)
        elif tag in ("em", "i"):
            _inline(paragraph, child, bold=bold,    italic=True,      underline=underline, strike=strike, code=code)
        elif tag == "u":
            _inline(paragraph, child, bold=bold,    italic=italic,    underline=True,      strike=strike, code=code)
        elif tag in ("s", "del", "strike"):
            _inline(paragraph, child, bold=bold,    italic=italic,    underline=underline, strike=True,   code=code)
        elif tag == "code":
            _inline(paragraph, child, bold=bold,    italic=italic,    underline=underline, strike=strike, code=True)
        elif tag == "a":
            href = child.get("href", "")
            if href:
                _add_hyperlink(paragraph, child.get_text(), href)
            else:
                _add_run(paragraph, child.get_text(), bold, italic, underline, strike, code)
        elif tag == "br":
            paragraph.add_run().add_break()
        elif tag == "img":
            src = child.get("src", "")
            img_bytes = _fetch_image(src) if src else None
            if img_bytes:
                try:
                    paragraph.add_run().add_picture(BytesIO(img_bytes), width=Inches(3))
                    continue
                except Exception:
                    pass
            alt = child.get("alt", "")
            _add_run(paragraph, f"[Image: {alt}]" if alt else "[Image]")
        elif tag == "span":
            _inline(paragraph, child, bold, italic, underline, strike, code)
        else:
            t = child.get_text()
            if t:
                _add_run(paragraph, t, bold, italic, underline, strike, code)


# ── Document processor ────────────────────────────────────────────────────────

_MD_EXTENSIONS = [
    "tables", "fenced_code", "nl2br", "sane_lists",
    "attr_list", "def_list", "footnotes", "md_in_html",
]

# MAIA template style mapping
# Keys are the markdown construct; values are tried in order (case-insensitive)
_STYLES = {
    "h1":         ["heading 1",   "Heading 1",   "Normal"],
    "h2":         ["heading 2",   "Heading 2",   "Normal"],
    "h3":         ["heading 3",   "Heading 3",   "Normal"],
    "h4":         ["heading 4",   "Heading 4",   "Normal"],
    "h5":         ["heading 5",   "Heading 5",   "Normal"],
    "h6":         ["heading 6",   "Heading 6",   "Normal"],
    "body":       ["Body Text",   "Normal"],
    "quote":      ["Anotacion",   "Intense Quote", "Quote", "Body Text", "Normal"],
    "code":       ["Normal"],
    "bullet":     ["List Bullet", "List Paragraph", "Normal"],
    "number":     ["List Number", "List Bullet",    "List Paragraph", "Normal"],
    "bullet_2":   ["List Bullet 2", "List Bullet",  "List Paragraph", "Normal"],
    "number_2":   ["List Number 2", "List Number",  "List Paragraph", "Normal"],
    "table":      ["Table Grid",  "Normal Table"],
}


class DocumentProcessor:
    """
    Usage:
        processor = DocumentProcessor()
        docx_bytes = processor.process(template_bytes, [md1, md2, md3, md4])
    """

    def __init__(self) -> None:
        self.doc: Optional[Document] = None
        self.styles: dict = {}

    def process(self, template_bytes: bytes, sections: list,
                fields: Optional[dict] = None) -> bytes:
        template_bytes = _normalize_docx_bytes(template_bytes)
        self.doc = Document(BytesIO(template_bytes))
        self.styles = _collect_styles(self.doc)
        if fields:
            self._fill_cover(fields)
        self._clear_body()
        for idx, md_text in enumerate(sections):
            if idx > 0:
                self._add_page_break()
            html = markdown.markdown(md_text, extensions=_MD_EXTENSIONS)
            soup = BeautifulSoup(f"<root>{html}</root>", "html.parser")
            self._process_children(soup.find("root"))
        out = BytesIO()
        self.doc.save(out)
        return out.getvalue()

    def _s(self, key: str) -> Optional[str]:
        """Resolve a style key against the live template styles."""
        return _best_style(self.styles, *_STYLES[key])

    def _fill_cover(self, fields: dict) -> None:
        """
        Fill all cover-page and running-header/footer placeholders from *fields*.

        Accepted keys (all optional — missing keys leave the placeholder intact):
          audit_code    e.g. "2601-0042"
          audit_title   e.g. "Auditoría de Controles TI"
          uai           e.g. "Tecnología"  → "UAI de Tecnología"
          date          e.g. "15/04/2026"  → replaces xx/xx/2026 everywhere
          recipients    list of strings    → fills D. / Dª. slots on cover
          audit_status  "BORRADOR" | "DEFINITIVO"  → replaces BORRADOR in header
        """
        audit_code   = (fields.get("audit_code")   or "").strip()
        audit_title  = (fields.get("audit_title")  or "").strip()
        uai          = (fields.get("uai")          or "").strip()
        date         = (fields.get("date")         or "").strip()
        recipients   = [r for r in (fields.get("recipients") or []) if r]
        audit_status = (fields.get("audit_status") or "").strip().upper()

        audit_label = (
            f"[{audit_code} {audit_title}]" if audit_code or audit_title else ""
        )
        uai_full = f"UAI de {uai}" if uai else ""

        # ── Cover body paragraphs ─────────────────────────────────────────────
        # Use .iter() to reach paragraphs inside text boxes (wps:txbxContent).
        # _fill_cover runs before _clear_body, so all template content is still
        # present — the placeholders are unique to the cover page.
        body = self.doc.element.body
        recipient_paras = []
        for p_elem in body.iter(qn("w:p")):
            text = _para_text(p_elem).strip()
            if audit_label and ("[26XX-XXXX" in text or "TÍTULO DE LA AUDITORÍA" in text):
                _set_para_text(p_elem, audit_label)
            elif uai_full and text == "UAI de XXX":
                _set_para_text(p_elem, uai_full)
            elif text in ("D.", "Dª."):
                recipient_paras.append(p_elem)

        if recipients:
            for i, p_elem in enumerate(recipient_paras):
                _set_para_text(p_elem, recipients[i % len(recipients)])

        # ── Headers and footers ───────────────────────────────────────────────
        for rel in self.doc.part.rels.values():
            rt = rel.reltype.lower()
            if "header" not in rt and "footer" not in rt:
                continue
            try:
                root = rel.target_part._element
            except Exception:
                continue

            for p_elem in root.iter(qn("w:p")):
                ptext = _para_text(p_elem)

                # Date: xx/xx/20xx pattern split across runs
                if date and re.search(r"xx/xx/20\d{2}", ptext, re.IGNORECASE):
                    _replace_date_runs(p_elem, date)

                # Header 2/3: UAI + title
                if uai and audit_title:
                    _replace_run_text(
                        p_elem,
                        "XXX \u2013 [T\u00edtulo de la Auditor\u00eda]",
                        f"{uai} \u2013 {audit_title}",
                    )

                # Header 1: BORRADOR → DEFINITIVO (only if explicitly requested)
                if audit_status and audit_status in ("BORRADOR", "DEFINITIVO"):
                    _replace_run_text(p_elem, "BORRADOR", audit_status)

    def _clear_body(self) -> None:
        """
        Preserve the cover page (first OOXML section) and remove everything
        after it.  The cover page ends with a <w:p> whose <w:pPr> contains a
        <w:sectPr> — that paragraph and everything before it is kept intact.
        Everything between it and the final body <w:sectPr> is removed so the
        4 report sections can be appended cleanly.
        """
        body = self.doc.element.body

        # Find the paragraph that closes the cover-page section
        cover_end = None
        for child in body:
            if child.tag == qn("w:p"):
                pPr = child.find(qn("w:pPr"))
                if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
                    cover_end = child
                    break

        if cover_end is not None:
            # Remove every body child AFTER the cover-page paragraph,
            # except the final <w:sectPr> (document-level section props)
            past_cover = False
            for elem in list(body):
                if elem is cover_end:
                    past_cover = True
                    continue
                if past_cover and elem.tag != qn("w:sectPr"):
                    body.remove(elem)
        else:
            # Fallback: no cover page — clear everything
            for elem in [c for c in body if c.tag != qn("w:sectPr")]:
                body.remove(elem)

        # Purge stale bookmark nodes to prevent ID collisions
        for bm_tag in (qn("w:bookmarkStart"), qn("w:bookmarkEnd")):
            for node in body.findall(f".//{bm_tag}"):
                parent = node.getparent()
                if parent is not None:
                    parent.remove(node)

    def _add_page_break(self) -> None:
        self.doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    def _process_children(self, element) -> None:
        for child in element.children:
            if isinstance(child, NavigableString):
                text = str(child).strip()
                if text:
                    self.doc.add_paragraph(text, style=self._s("body"))
            elif hasattr(child, "name") and child.name:
                self._process_block(child)

    def _process_block(self, element) -> None:
        tag = element.name.lower()

        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            p = self.doc.add_paragraph(style=self._s(tag))
            _inline(p, element)

        elif tag == "p":
            imgs = element.find_all("img")
            if imgs and not element.get_text(strip=True):
                p = self.doc.add_paragraph()
                for img in imgs:
                    img_bytes = _fetch_image(img.get("src", ""))
                    if img_bytes:
                        try:
                            p.add_run().add_picture(BytesIO(img_bytes), width=Inches(4))
                            continue
                        except Exception:
                            pass
                    alt = img.get("alt", "")
                    p.add_run(f"[Image: {alt}]" if alt else "[Image]")
            else:
                p = self.doc.add_paragraph(style=self._s("body"))
                _inline(p, element)

        elif tag == "ul":
            self._process_list(element, ordered=False, level=0)
        elif tag == "ol":
            self._process_list(element, ordered=True, level=0)

        elif tag == "blockquote":
            for child in element.children:
                if isinstance(child, NavigableString):
                    text = str(child).strip()
                    if text:
                        self.doc.add_paragraph(text, style=self._s("quote"))
                elif hasattr(child, "name") and child.name in ("p", "div", "span"):
                    p = self.doc.add_paragraph(style=self._s("quote"))
                    _inline(p, child)
                else:
                    self._process_block(child)

        elif tag == "pre":
            code_elem = element.find("code")
            code_text = (code_elem if code_elem else element).get_text()
            p = self.doc.add_paragraph(style=self._s("code"))
            run = p.add_run(code_text)
            run.font.name = "Courier New"
            run.font.size = Pt(9)

        elif tag == "table":
            self._process_table(element)

        elif tag == "hr":
            _add_horizontal_rule(self.doc.add_paragraph())

        elif tag == "img":
            p = self.doc.add_paragraph()
            img_bytes = _fetch_image(element.get("src", ""))
            if img_bytes:
                try:
                    p.add_run().add_picture(BytesIO(img_bytes), width=Inches(4))
                    return
                except Exception:
                    pass
            alt = element.get("alt", "")
            p.add_run(f"[Image: {alt}]" if alt else "[Image]")

        elif tag in ("div", "section", "article", "main", "aside",
                     "header", "footer", "figure", "figcaption"):
            self._process_children(element)

    def _process_list(self, element, ordered: bool, level: int) -> None:
        if ordered:
            style_key = "number" if level == 0 else "number_2"
        else:
            style_key = "bullet" if level == 0 else "bullet_2"

        for child in element.children:
            if not hasattr(child, "name") or child.name != "li":
                continue
            nested = child.find(["ul", "ol"])
            p = self.doc.add_paragraph(style=self._s(style_key))
            if level > 0:
                p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
            for item_child in child.children:
                if hasattr(item_child, "name") and item_child.name in ("ul", "ol"):
                    continue
                if isinstance(item_child, NavigableString):
                    if str(item_child).strip():
                        p.add_run(str(item_child))
                elif hasattr(item_child, "name"):
                    _inline(p, item_child)
            if nested:
                self._process_list(nested, ordered=(nested.name == "ol"), level=level + 1)

    def _process_table(self, element) -> None:
        rows = element.find_all("tr")
        if not rows:
            return
        max_cols = max((len(r.find_all(["td", "th"])) for r in rows), default=0)
        if not max_cols:
            return
        table = self.doc.add_table(rows=len(rows), cols=max_cols)
        tbl_style = self._s("table")
        if tbl_style:
            try:
                table.style = tbl_style
            except Exception:
                pass
        for row_idx, tr in enumerate(rows):
            for col_idx, cell_elem in enumerate(tr.find_all(["td", "th"])):
                if col_idx >= max_cols:
                    break
                cell = table.cell(row_idx, col_idx)
                cell.paragraphs[0].clear()
                p = cell.paragraphs[0]
                _inline(p, cell_elem)
                if cell_elem.name == "th":
                    for run in p.runs:
                        run.bold = True
