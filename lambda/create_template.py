"""
Generates the corporate .docx template and uploads it to S3.

Usage:
  python create_template.py                  # saves locally to templates/corporate_template.docx
  python create_template.py --upload         # also uploads to S3 (reads .env)
"""

import argparse
import os
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

OUTPUT_PATH = Path(__file__).parent.parent / "templates" / "corporate_template.docx"

# PwC brand colours
_RED = RGBColor(0xD0, 0x43, 0x3B)
_DARK = RGBColor(0x21, 0x21, 0x21)
_GREY = RGBColor(0x76, 0x76, 0x76)


def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def build_template() -> bytes:
    doc = Document()

    # ------------------------------------------------------------------ #
    # Page layout — A4, narrow margins
    # ------------------------------------------------------------------ #
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # ------------------------------------------------------------------ #
    # Styles
    # ------------------------------------------------------------------ #
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = _DARK

    for level in range(1, 5):
        h = styles[f"Heading {level}"]
        h.font.name = "Calibri"
        h.font.color.rgb = _RED if level == 1 else _DARK
        h.font.bold = True
        h.font.size = Pt(20 - (level - 1) * 2)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)

    body_text = styles.add_style("Body Text", 1)
    body_text.base_style = styles["Normal"]
    body_text.font.size = Pt(11)
    body_text.paragraph_format.space_after = Pt(8)

    # ------------------------------------------------------------------ #
    # Header — red accent bar + company name
    # ------------------------------------------------------------------ #
    header = section.header
    header.is_linked_to_previous = False

    # Red top bar via a 1-row, 1-col table
    tbl = header.add_table(rows=1, cols=1, width=Cm(16))
    tbl.style = "Table Normal"
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, "D0433B")
    cell.paragraphs[0].clear()
    run = cell.paragraphs[0].add_run("  PwC")
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.name = "Calibri"

    # Company subtitle below the bar
    p = header.add_paragraph("PricewaterhouseCoopers  |  Confidential")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.runs[0]
    run.font.size = Pt(8)
    run.font.color.rgb = _GREY
    run.font.name = "Calibri"

    # ------------------------------------------------------------------ #
    # Footer — page numbers
    # ------------------------------------------------------------------ #
    footer = section.footer
    footer.is_linked_to_previous = False

    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.clear()

    run = fp.add_run("Page ")
    run.font.size = Pt(9)
    run.font.color.rgb = _GREY

    # PAGE field
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    run._r.append(fld)
    instr = OxmlElement("w:instrText")
    instr.text = " PAGE "
    run._r.append(instr)
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld2)

    run2 = fp.add_run(" of ")
    run2.font.size = Pt(9)
    run2.font.color.rgb = _GREY

    # NUMPAGES field
    fld3 = OxmlElement("w:fldChar")
    fld3.set(qn("w:fldCharType"), "begin")
    run2._r.append(fld3)
    instr2 = OxmlElement("w:instrText")
    instr2.text = " NUMPAGES "
    run2._r.append(instr2)
    fld4 = OxmlElement("w:fldChar")
    fld4.set(qn("w:fldCharType"), "end")
    run2._r.append(fld4)

    # ------------------------------------------------------------------ #
    # Placeholder body paragraph (will be replaced by the Lambda)
    # ------------------------------------------------------------------ #
    p = doc.add_paragraph()
    p.style = styles["Normal"]

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def upload_to_s3(template_bytes: bytes) -> None:
    import boto3
    bucket = os.environ["S3_BUCKET_NAME"]
    key = os.environ.get("TEMPLATE_S3_KEY", "templates/corporate_template.docx")
    s3 = boto3.client("s3", region_name=os.environ.get("AWS_REGION", "us-east-1"))
    s3.put_object(
        Bucket=bucket,
        Key=key,
        Body=template_bytes,
        ContentType=(
            "application/vnd.openxmlformats-officedocument"
            ".wordprocessingml.document"
        ),
    )
    print(f"Uploaded to s3://{bucket}/{key}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--upload", action="store_true", help="Upload template to S3")
    args = parser.parse_args()

    template_bytes = build_template()

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT_PATH.write_bytes(template_bytes)
    print(f"Saved to {OUTPUT_PATH}")

    if args.upload:
        upload_to_s3(template_bytes)
