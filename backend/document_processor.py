"""
Core document processor: loads a Word template, clears its body,
then rebuilds content from parsed markdown — preserving all template
styles, page layout, headers, footers, and images.
"""

import base64
import re
import urllib.request
from io import BytesIO
import os
import subprocess  # Not strictly needed for this approach, but good to keep if Pandoc is auxiliary
import tempfile  # Not strictly needed for this approach, but good to keep if Pandoc is auxiliary
from typing import Optional, Literal, Dict, Any

import markdown
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
import yaml


# ---------------------------------------------------------------------------
# Mapfre font map — derived from the template's styles.xml definitions.
# Keys are OOXML style IDs (w:styleId); value is (ascii, hAnsi, cs).
# None covers paragraphs with no explicit pStyle (defaults to Normal).
# ---------------------------------------------------------------------------

_DISPLAY = "Mapfre Display"
_TEXT = "Mapfre Text"

# Headings 1-3: Display for all three font slots (as in styles.xml)
# Headings 4-9: Display for ascii/hAnsi; cs follows the template (Times New Roman)
# Body, lists, Normal, fallback: Text for all slots
_MAPFRE_RFONTS: dict[Optional[str], tuple[str, str, str]] = {
    "Heading1": (_DISPLAY, _DISPLAY, _DISPLAY),
    "Heading2": (_DISPLAY, _DISPLAY, _DISPLAY),
    "Heading3": (_DISPLAY, _DISPLAY, _DISPLAY),
    "Heading4": (_DISPLAY, _DISPLAY, "Times New Roman"),
    "Heading5": (_DISPLAY, _DISPLAY, "Times New Roman"),
    "Heading6": (_DISPLAY, _DISPLAY, "Times New Roman"),
    "Heading7": (_DISPLAY, _DISPLAY, "Times New Roman"),
    "Heading8": (_DISPLAY, _DISPLAY, "Times New Roman"),
    "Heading9": (_DISPLAY, _DISPLAY, "Times New Roman"),
    "Caption": (_DISPLAY, _DISPLAY, "Times New Roman"),
    "Cabecera1": (_DISPLAY, _DISPLAY, "Times New Roman"),
    "Cabecera2": (_DISPLAY, _DISPLAY, "Times New Roman"),
    "Anexo1": (_DISPLAY, _DISPLAY, "Times New Roman"),
    "Anexo2": (_DISPLAY, _DISPLAY, "Times New Roman"),
    "BodyText": (_TEXT, _TEXT, _TEXT),
    None: (_TEXT, _TEXT, _TEXT),  # Normal (no explicit pStyle)
    "Normal": (_TEXT, _TEXT, _TEXT),
    "ListParagraph": (_TEXT, _TEXT, _TEXT),
    "ListBullet": (_TEXT, _TEXT, _TEXT),
    "ListBullet2": (_TEXT, _TEXT, _TEXT),
    "ListBullet3": (_TEXT, _TEXT, _TEXT),
    "ListNumber": (_TEXT, _TEXT, _TEXT),
    "ListNumber2": (_TEXT, _TEXT, _TEXT),
    "ListNumber3": (_TEXT, _TEXT, _TEXT),
    "IntenseQuote": (_TEXT, _TEXT, _TEXT),
    "Quote": (_TEXT, _TEXT, _TEXT),
    "NoSpacing": (_TEXT, _TEXT, _TEXT),
}


def _make_rfonts(ascii_f: str, hansi_f: str, cs_f: str):
    """Create a fresh <w:rFonts> element with the given font slots."""
    elem = OxmlElement("w:rFonts")
    elem.set(qn("w:ascii"), ascii_f)
    elem.set(qn("w:hAnsi"), hansi_f)
    elem.set(qn("w:cs"), cs_f)
    return elem


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _collect_style_names(doc: Document) -> dict[str, set]:
    paragraph_styles = {s.name for s in doc.styles if s.type == WD_STYLE_TYPE.PARAGRAPH}
    character_styles = {s.name for s in doc.styles if s.type == WD_STYLE_TYPE.CHARACTER}
    return {"paragraph": paragraph_styles, "character": character_styles}


def _best_style(
    available_styles: dict[str, set], style_type: str, *candidates: str
) -> Optional[str]:
    """Return the first candidate of a given type that exists in the document styles."""
    if style_type not in available_styles:
        return None
    styles_of_type = available_styles[style_type]
    for name in candidates:
        if name and name in styles_of_type:
            return name
    return None


def _add_horizontal_rule(paragraph) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_hyperlink(
    document: Document,
    available_char_styles: set,
    paragraph,
    text: str,
    url: str,
    config_char_styles: Dict[str, str],  # Added config_char_styles
) -> None:
    if not url:
        paragraph.add_run(text)
        return
    try:
        r_id = paragraph.part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        new_run = OxmlElement("w:r")

        # Try to apply 'Hyperlink' character style from config or template
        mapped_style = config_char_styles.get("a")
        hyperlink_style = _best_style(
            {"character": available_char_styles}, "character", mapped_style, "Hyperlink"
        )

        if hyperlink_style:
            rPr = OxmlElement("w:rPr")
            rStyle = OxmlElement("w:rStyle")
            rStyle.set(qn("w:val"), hyperlink_style)
            rPr.append(rStyle)
            new_run.append(rPr)
        else:
            # Fallback to direct underline if no 'Hyperlink' style mapped or found
            rPr = OxmlElement("w:rPr")
            u = OxmlElement("w:u")
            u.set(qn("w:val"), "single")
            rPr.append(u)
            new_run.append(rPr)

        t = OxmlElement("w:t")
        t.text = text
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        new_run.append(t)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
    except Exception:
        run = paragraph.add_run(text)
        run.underline = True


def _fetch_image(src: str) -> Optional[bytes]:
    """Return raw bytes for a data-URI or http(s) image URL."""
    if src.startswith("data:"):
        m = re.match(r"data:[^;]+;base64,(.+)", src, re.DOTALL)
        if m:
            return base64.b64decode(m.group(1))
    elif src.startswith(("http://", "https://")):
        try:
            req = urllib.request.Request(src, headers={"User-Agent": "Mozilla/5.0"})
            with urllib.request.urlopen(req, timeout=5) as resp:
                return resp.read()
        except Exception:
            return None
    return None


# ---------------------------------------------------------------------------
# Inline content builder (for python-docx manual approach) - Refined for config
# ---------------------------------------------------------------------------


def _add_run_with_fmt(
    document: Document,
    available_char_styles: set,
    paragraph,
    text: str,
    char_style_name: Optional[str] = None,
) -> None:
    if not text:
        return

    run = paragraph.add_run(text)

    # Prioritize explicit character style if provided and it exists in available styles
    if char_style_name and char_style_name in available_char_styles:
        run.style = document.styles[char_style_name]
    # Direct formatting fallbacks for semantic elements if no matching character style
    elif char_style_name == "Strong":
        run.bold = True
    elif char_style_name == "Emphasis":
        run.italic = True
    elif char_style_name == "Underline":
        run.underline = True
    elif char_style_name == "Strikethrough":
        run.font.strike = True
    elif char_style_name == "Code":
        run.font.name = "Courier New"
        run.font.size = Pt(9)


def _inline(
    document: Document,
    available_char_styles: set,
    paragraph,
    element,
    config_char_styles: Dict[str, str],  # Added config_char_styles
    inherited_char_style_name: Optional[str] = None,
) -> None:
    """Recursively walk inline HTML elements and add runs to *paragraph*."""
    for child in element.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if text:
                _add_run_with_fmt(
                    document,
                    available_char_styles,
                    paragraph,
                    text,
                    inherited_char_style_name,
                )
            continue

        if not hasattr(child, "name") or child.name is None:
            continue

        tag = child.name.lower()

        current_char_style_name = (
            inherited_char_style_name  # Start with inherited style
        )

        # Determine specific character style for known inline tags
        # Prioritize config.yaml mapping, then best_style from template, then semantic hint for direct fallback
        if tag in ("strong", "b"):
            mapped_style = config_char_styles.get("strong")
            current_char_style_name = (
                _best_style(
                    {"character": available_char_styles},
                    "character",
                    mapped_style,  # Use mapped style from config first
                    "Strong",
                    "Emphasis",
                )
                or "Strong"  # Conceptual name for direct formatting fallback in _add_run_with_fmt
            )
        elif tag in ("em", "i"):
            mapped_style = config_char_styles.get("em")
            current_char_style_name = (
                _best_style(
                    {"character": available_char_styles},
                    "character",
                    mapped_style,
                    "Emphasis",
                )
                or "Emphasis"  # Conceptual name for direct formatting fallback
            )
        elif tag == "u":
            mapped_style = config_char_styles.get("u")
            current_char_style_name = (
                _best_style(
                    {"character": available_char_styles},
                    "character",
                    mapped_style,
                    "Underline",
                )
                or "Underline"  # Conceptual name for direct formatting fallback
            )
        elif tag in ("s", "del", "strike"):
            mapped_style = config_char_styles.get("s")
            current_char_style_name = (
                _best_style(
                    {"character": available_char_styles},
                    "character",
                    mapped_style,
                    "Strikethrough",
                )
                or "Strikethrough"  # Conceptual name for direct formatting fallback
            )
        elif tag == "code":
            mapped_style = config_char_styles.get("code")
            current_char_style_name = (
                _best_style(
                    {"character": available_char_styles},
                    "character",
                    mapped_style,
                    "Code",
                    "Literal",
                )
                or "Code"  # Conceptual name for direct formatting fallback
            )
        elif tag == "a":
            href = child.get("href", "")
            link_text = child.get_text()
            if href:
                _add_hyperlink(
                    document,
                    available_char_styles,
                    paragraph,
                    link_text,
                    href,
                    config_char_styles,
                )
            else:
                _add_run_with_fmt(
                    document,
                    available_char_styles,
                    paragraph,
                    link_text,
                    current_char_style_name,
                )
            # Hyperlinks are self-contained, no need to recurse for text children in <a>
            continue
        elif tag == "br":
            paragraph.add_run("\n")
            continue
        elif tag == "img":
            src = child.get("src", "")
            alt = child.get("alt", "")
            if src:
                img_bytes = _fetch_image(src)
                if img_bytes:
                    try:
                        run = paragraph.add_run()
                        run.add_picture(BytesIO(img_bytes), width=Inches(3))
                    except Exception:
                        _add_run_with_fmt(
                            document,
                            available_char_styles,
                            paragraph,
                            f"[Image: {alt}]",
                            current_char_style_name,
                        )
                else:
                    _add_run_with_fmt(
                        document,
                        available_char_styles,
                        paragraph,
                        f"[Image: {alt}]",
                        current_char_style_name,
                    )
            elif alt:
                _add_run_with_fmt(
                    document,
                    available_char_styles,
                    paragraph,
                    f"[Image: {alt}]",
                    current_char_style_name,
                )
            continue  # Image is self-contained

        # For other tags (like <span> or if a style was found), recurse with the new or inherited style
        _inline(
            document,
            available_char_styles,
            paragraph,
            child,
            config_char_styles,
            current_char_style_name,
        )


# ---------------------------------------------------------------------------
# Configuration Loader
# ---------------------------------------------------------------------------


class ConfigLoader:
    def __init__(self, config_path: str):
        self.config_path = config_path
        self.config: Dict[str, Any] = self._load_config()

    def _load_config(self) -> Dict[str, Any]:
        if not os.path.exists(self.config_path):
            # Return default empty config if file doesn't exist
            return {
                "paragraph_styles": {},
                "character_styles": {},
                "list_numbering": {},
            }
        with open(self.config_path, "r", encoding="utf-8") as f:
            return yaml.safe_load(f)

    def get_paragraph_style(
        self, markdown_tag: str, default: Optional[str] = None
    ) -> Optional[str]:
        return self.config.get("paragraph_styles", {}).get(markdown_tag, default)

    def get_character_style(
        self, markdown_tag: str, default: Optional[str] = None
    ) -> Optional[str]:
        return self.config.get("character_styles", {}).get(markdown_tag, default)

    def get_list_numbering_config(self) -> Dict[str, str]:
        return self.config.get("list_numbering", {})

    def get_content_placeholder(self) -> str:
        return self.config.get("content_placeholder", "{{markdown_content}}")


# ---------------------------------------------------------------------------
# Main processor class
# ---------------------------------------------------------------------------


class DocumentProcessor:
    """
    Loads a .docx template, clears its body,
    then rebuilds content from parsed markdown — preserving all template
    styles, page layout, headers, footers, and images.
    """

    def __init__(
        self, config_path: str = "/home/diego/dev/pwc/backend/styles_map.yaml"
    ):
        self.doc: Optional[Document] = None
        self.styles: dict[str, set] = {"paragraph": set(), "character": set()}
        self.config_loader = ConfigLoader(config_path)

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def process(
        self,
        template_docx_path: str,
        markdown_text_sections: Dict[
            str, str
        ],  # Now takes a dictionary of markdown sections
        cover_fields: Optional[dict] = None,
    ) -> bytes:
        return self.process_with_targeted_injection(
            template_docx_path, markdown_text_sections, cover_fields
        )

    def _add_paragraph_from_html_element(
        self,
        parent_body,
        html_element,
        placeholder_idx,
        config_para_styles,
        config_char_styles,
        available_char_styles,
    ):
        # Helper to convert HTML elements to docx paragraphs and insert
        temp_doc_for_block = Document()
        temp_styles = _collect_style_names(temp_doc_for_block)
        temp_available_char_styles = temp_styles["character"]

        # Temporarily redirect self.doc and self.styles to build content into temp_doc_for_block
        original_doc = self.doc
        original_styles = self.styles
        self.doc = temp_doc_for_block
        self.styles = temp_styles

        self._process_block(
            html_element, self.doc, temp_available_char_styles, config_char_styles
        )

        self.doc = original_doc  # Switch back
        self.styles = original_styles

        for elem_to_insert in reversed(temp_doc_for_block.element.body):
            parent_body.insert(placeholder_idx, elem_to_insert)

    def process_with_targeted_injection(
        self,
        template_docx_path: str,
        markdown_text_sections: Dict[
            str, str
        ],  # Dictionary of {placeholder: markdown_content}
        cover_fields: Optional[dict] = None,
    ) -> bytes:
        # Step 1: Load the template DOCX with python-docx
        self.doc = Document(template_docx_path)
        self.styles = _collect_style_names(self.doc)
        config_para_styles_map = self.config_loader.config.get("paragraph_styles", {})
        config_char_styles_map = self.config_loader.config.get("character_styles", {})
        list_numbering_config = self.config_loader.get_list_numbering_config()

        available_char_styles = self.styles["character"]

        # Store original template_bytes for potential use in Pandoc reference
        original_template_bytes = None
        try:
            with open("/home/diego/dev/pwc/backend/template_base64.txt", "r") as f:
                original_template_base64 = f.read().strip()
            original_template_bytes = base64.b64decode(original_template_base64)
        except FileNotFoundError:
            raise FileNotFoundError(
                "template_base64.txt not found. Please ensure it exists."
            )

        temp_original_docx_path = ""
        try:
            with tempfile.NamedTemporaryFile(
                suffix=".docx", delete=False
            ) as temp_orig_template_file:
                temp_orig_template_file.write(original_template_bytes)
                temp_original_docx_path = temp_orig_template_file.name

            # Step 2 & 3: Iterate through placeholders and inject content
            for placeholder_text, markdown_text in markdown_text_sections.items():
                found_placeholder = False
                for paragraph in self.doc.paragraphs:
                    if placeholder_text in paragraph.text:
                        parent_body = paragraph._element.getparent()
                        if parent_body is None:
                            continue

                        placeholder_idx = parent_body.index(paragraph._element)
                        parent_body.remove(paragraph._element)  # Remove placeholder
                        found_placeholder = True

                        # Convert markdown to DOCX fragment using Pandoc, referencing the original template
                        pandoc_content_docx_path = os.path.join(
                            tempfile.gettempdir(),
                            f"pandoc_content_{os.urandom(8).hex()}.docx",
                        )
                        temp_markdown_file_path = os.path.join(
                            tempfile.gettempdir(),
                            f"temp_markdown_{os.urandom(8).hex()}.md",
                        )

                        try:
                            with open(
                                temp_markdown_file_path, "w", encoding="utf-8"
                            ) as f:
                                f.write(markdown_text)
                            command = [
                                "pandoc",
                                "-s",
                                temp_markdown_file_path,
                                "-o",
                                pandoc_content_docx_path,
                                "--reference-doc",
                                temp_original_docx_path,  # Use original template as reference
                            ]
                            subprocess.run(command, check=True, capture_output=True)

                            # Load Pandoc-generated content and inject
                            pandoc_content_doc = Document(pandoc_content_docx_path)
                            for content_element in reversed(
                                pandoc_content_doc.element.body
                            ):
                                parent_body.insert(placeholder_idx, content_element)

                        finally:
                            if os.path.exists(temp_markdown_file_path):
                                os.remove(temp_markdown_file_path)
                            if os.path.exists(pandoc_content_docx_path):
                                os.remove(pandoc_content_docx_path)
                        break  # Placeholder processed, move to next

                if not found_placeholder:
                    print(
                        f"Warning: Placeholder '{placeholder_text}' not found in template. Content not inserted."
                    )

            # Step 4: Apply cover fields using python-docx
            if cover_fields:
                self._fill_cover_page(cover_fields)

            # Step 5: Post-process lists to ensure they use template numbering schemes
            self._remap_list_styles(list_numbering_config, config_para_styles_map)

            # Step 6: Stamp fonts for consistency
            self._stamp_run_fonts()

            final_output_buffer = BytesIO()
            self.doc.save(final_output_buffer)
            return final_output_buffer.getvalue()

        finally:
            if os.path.exists(temp_original_docx_path):
                os.remove(temp_original_docx_path)

    def _find_cover_end(self) -> int:
        """Return the index of the cover-page section-break paragraph."""
        body = self.doc.element.body
        for i, child in enumerate(body):
            if child.tag == qn("w:p"):
                pPr = child.find(qn("w:pPr"))
                if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
                    return i
        return 0

    def _stamp_run_fonts(self) -> None:
        """
        Stamp the hardcoded Mapfre rFonts (from _MAPFRE_RFONTS) onto every
        generated run that doesn't already carry an explicit <w:rFonts>.
        Runs inside the cover page are skipped.
        Code runs (Courier New) are also skipped.
        """
        body = self.doc.element.body
        children = list(body)
        cover_end = self._find_cover_end()

        for child in children[cover_end + 1 :]:
            if child.tag != qn("w:p"):
                continue

            pPr = child.find(qn("w:pPr"))
            pStyle_elem = pPr.find(qn("w:pStyle")) if pPr is not None else None
            style_id = pStyle_elem.get(qn("w:val")) if pStyle_elem is not None else None

            fonts = _MAPFRE_RFONTS.get(style_id, _MAPFRE_RFONTS[None])

            for run in child.findall(qn("w:r")):
                rPr = run.find(qn("w:rPr"))
                if rPr is None:
                    rPr = OxmlElement("w:rPr")
                    run.insert(0, rPr)
                if rPr.find(qn("w:rFonts")) is None:
                    rPr.insert(0, _make_rfonts(*fonts))

    def _remap_list_styles(
        self,
        list_numbering_config: Dict[str, str],
        config_para_styles_map: Dict[str, str],
    ) -> None:
        """Post-process paragraphs to remap Pandoc's list numIds to template-defined ones."""
        bullet_num_id = list_numbering_config.get("bullet_num_id")
        decimal_num_id = list_numbering_config.get("decimal_num_id")

        if not bullet_num_id or not decimal_num_id:
            print(
                "Warning: List numbering IDs not found in styles_map.yaml. List remapping skipped."
            )
            return

        for paragraph in self.doc.paragraphs:
            p_pr = paragraph._p.pPr
            if p_pr is not None:
                num_pr = p_pr.numPr
                if num_pr is not None:
                    ilvl_elem = num_pr.ilvl
                    numId_elem = num_pr.numId

                    if ilvl_elem is not None and numId_elem is not None:
                        current_ilvl = ilvl_elem.val
                        current_numId = numId_elem.val

                        # Get style hints from config and paragraph's actual style
                        p_style_elem = p_pr.pStyle
                        current_p_style_name = (
                            p_style_elem.val if p_style_elem is not None else None
                        )

                        mapped_bullet_para_style = config_para_styles_map.get(
                            "list_item_bullet"
                        )
                        mapped_number_para_style = config_para_styles_map.get(
                            "list_item_number"
                        )

                        # Determine if this paragraph should be a bullet or numbered list item
                        should_be_bullet = False
                        should_be_number = False

                        if (
                            current_p_style_name == mapped_bullet_para_style
                            or "List Bullet" in current_p_style_name
                            or "List Paragraph" in current_p_style_name
                        ):
                            should_be_bullet = True
                        elif (
                            current_p_style_name == mapped_number_para_style
                            or "List Number" in current_p_style_name
                        ):
                            should_be_number = True

                        if should_be_bullet and current_numId != bullet_num_id:
                            num_pr.numId.val = bullet_num_id
                            num_pr.ilvl.val = current_ilvl
                        elif should_be_number and current_numId != decimal_num_id:
                            num_pr.numId.val = decimal_num_id
                            num_pr.ilvl.val = current_ilvl
                        elif current_numId not in [bullet_num_id, decimal_num_id]:
                            # Fallback if no clear style hint and not already mapped
                            if current_ilvl == "0":
                                num_pr.numId.val = bullet_num_id
                            else:
                                num_pr.numId.val = decimal_num_id
                            num_pr.ilvl.val = current_ilvl

    @staticmethod
    def _replace_across_runs(element, find: str, replace: str) -> None:
        """Replace *find* with *replace* even when the text spans multiple w:t nodes."""
        while True:
            t_nodes = list(element.iter(qn("w:t")))
            texts = [t.text or "" for t in t_nodes]
            full = "".join(texts)
            if find not in full:
                break

            start = full.index(find)
            end = start + len(find)

            # Map character ranges to node indices
            cum, spans = 0, []
            for txt in texts:
                spans.append((cum, cum + len(txt)))
                cum += len(txt)

            affected = [
                i for i, (ns, ne) in enumerate(spans) if ne > start and ns < end
            ]
            if not affected:
                break

            fi, li = affected[0], affected[-1]
            prefix = texts[fi][: max(0, start - spans[fi][0])]
            suffix = texts[li][max(0, end - spans[li][0]) :]

            if fi == li:
                t_nodes[fi].text = prefix + replace + suffix
            else:
                t_nodes[fi].text = prefix + replace
                for i in affected[1:-1]:
                    t_nodes[i].text = ""
                t_nodes[li].text = suffix

    def _fill_cover_page(self, cover_fields: dict) -> None:
        """Replace cover-page placeholders with the values in *cover_fields*."""
        body = self.doc.element.body
        children = list(body)
        cover_end = self._find_cover_end()

        audit_code = cover_fields.get("audit_code", "")
        audit_title = cover_fields.get("audit_title", "")
        uai = cover_fields.get("uai", "")
        date = cover_fields.get("date", "")
        recipients = cover_fields.get("recipients", [])

        for elem in children[: cover_end + 1]:
            if audit_code:
                self._replace_across_runs(elem, "26XX-XXXX", audit_code)
            if audit_title:
                self._replace_across_runs(
                    elem, "TÍTULO DE LA AUDITORÍA", audit_title.upper()
                )
            if uai:
                self._replace_across_runs(elem, "UAI de XXX", uai)
            if date:
                for placeholder in ("xx/xx/xxxx", "XX/XX/XXXX", "xx/XX/XXXX"):
                    self._replace_across_runs(elem, placeholder, date)

        # Recipients — each `D. ` or `Dª. ` w:t node is a placeholder slot
        if recipients:
            slot_tags = ("D. ", "Dª. ")
            slot_idx = 0
            for elem in children[: cover_end + 1]:
                for t in elem.iter(qn("w:t")):
                    if slot_idx >= len(recipients):
                        break
                    if t.text in slot_tags:
                        t.text = recipients[slot_idx]
                        slot_idx += 1

    # --- Tree traversal (for python-docx manual approach) - Re-integrated for placeholder injection ---
    def _process_children(
        self,
        element,
        document: Document,
        available_char_styles: set,
        config_char_styles: Dict[str, str],
    ) -> None:
        for child in element.children:
            if isinstance(child, NavigableString):
                text = str(child).strip()
                if text:
                    self._add_paragraph(
                        text, self.config_loader.get_paragraph_style("p", "Normal")
                    )
            elif hasattr(child, "name") and child.name:
                self._process_block(
                    child, document, available_char_styles, config_char_styles
                )

    def _process_block(
        self,
        element,
        document: Document,
        available_char_styles: set,
        config_char_styles: Dict[str, str],
    ) -> None:
        tag = element.name.lower()
        config_para_styles = self.config_loader.config.get("paragraph_styles", {})

        # --- Headings ---
        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = int(tag[1])
            config_style = config_para_styles.get(tag)
            style = _best_style(
                self.styles["paragraph"],
                "paragraph",
                config_style,
                f"Heading {level}",
                "Normal",
            )
            p = self.doc.add_paragraph(style=style)
            _inline(
                document, available_char_styles, p, element, config_char_styles, None
            )

        # --- Paragraph ---
        elif tag == "p":
            imgs = element.find_all("img")
            if imgs and not element.get_text(strip=True):
                p = self.doc.add_paragraph()
                for img in imgs:
                    src = img.get("src", "")
                    alt = img.get("alt", "")
                    if src:
                        img_bytes = _fetch_image(src)
                        if img_bytes:
                            try:
                                run = p.add_run()
                                run.add_picture(BytesIO(img_bytes), width=Inches(4))
                                continue
                            except Exception:
                                pass
                    _add_run_with_fmt(
                        document, available_char_styles, p, f"[Image: {alt}]"
                    )
            else:
                config_style = config_para_styles.get("p")
                style = _best_style(
                    self.styles["paragraph"],
                    "paragraph",
                    config_style,
                    "Body Text",
                    "Normal",
                )
                p = self.doc.add_paragraph(style=style)
            _inline(
                document, available_char_styles, p, element, config_char_styles, None
            )

        # --- Lists ---
        elif tag == "ul":
            self._process_list(
                element,
                ordered=False,
                level=0,
                document=document,
                available_char_styles=available_char_styles,
                config_para_styles=config_para_styles,
                config_char_styles=config_char_styles,  # Pass config_char_styles
            )
        elif tag == "ol":
            self._process_list(
                element,
                ordered=True,
                level=0,
                document=document,
                available_char_styles=available_char_styles,
                config_para_styles=config_para_styles,
                config_char_styles=config_char_styles,  # Pass config_char_styles
            )

        # --- Blockquote ---
        elif tag == "blockquote":
            config_style = config_para_styles.get("blockquote")
            q_style = _best_style(
                self.styles["paragraph"],
                "paragraph",
                config_style,
                "Intense Quote",
                "Quote",
                "Body Text",
                "Normal",
            )
            for child in element.children:
                if isinstance(child, NavigableString):
                    text = str(child).strip()
                    if text:
                        self.doc.add_paragraph(text, style=q_style)
                elif hasattr(child, "name") and child.name in ("p", "div", "span"):
                    p = self.doc.add_paragraph(style=q_style)
                    _inline(
                        document,
                        available_char_styles,
                        p,
                        child,
                        config_char_styles,
                        None,
                    )
                else:
                    # Recurse for nested blockquotes or other blocks
                    self._process_block(
                        child, document, available_char_styles, config_char_styles
                    )

        # --- Code block ---
        elif tag == "pre":
            code_elem = element.find("code")
            code_text = (code_elem if code_elem else element).get_text()
            config_style = config_para_styles.get("code_block")
            style = _best_style(
                self.styles["paragraph"],
                "paragraph",
                config_style,
                "Code",
                "No Spacing",
                "Normal",
            )
            p = self.doc.add_paragraph(style=style)
            run = p.add_run(code_text)
            if style in (None, "Normal", "No Spacing"):
                run.font.name = "Courier New"
                run.font.size = Pt(9)

        # --- Table ---
        elif tag == "table":
            self._process_table(
                element, document, available_char_styles, config_char_styles
            )

        # --- Horizontal rule ---
        elif tag == "hr":
            p = self.doc.add_paragraph()
            _add_horizontal_rule(p)

        # --- Standalone image ---
        elif tag == "img":
            p = self.doc.add_paragraph()
            src = element.get("src", "")
            alt = element.get("alt", "")
            if src:
                img_bytes = _fetch_image(src)
                if img_bytes:
                    try:
                        run = p.add_run()
                        run.add_picture(BytesIO(img_bytes), width=Inches(4))
                        return
                    except Exception:
                        pass
            _add_run_with_fmt(
                document,
                available_char_styles,
                p,
                f"[Image: {alt}]" if alt else "[Image]",
            )

        # --- Generic container ---
        elif tag in (
            "div",
            "section",
            "article",
            "main",
            "aside",
            "header",
            "footer",
            "figure",
            "figcaption",
        ):
            self._process_children(
                element, document, available_char_styles, config_char_styles
            )

    def _process_list(
        self,
        element,
        ordered: bool,
        level: int,
        document: Document,
        available_char_styles: set,
        config_para_styles: Dict[str, str],  # Added config_para_styles
        config_char_styles: Dict[str, str],  # Added config_char_styles
    ) -> None:
        list_numbering_config = self.config_loader.get_list_numbering_config()
        bullet_num_id = list_numbering_config.get("bullet_num_id")
        decimal_num_id = list_numbering_config.get("decimal_num_id")

        if not bullet_num_id or not decimal_num_id:
            print(
                "Warning: List numbering IDs not found in styles_map.yaml. List styling might be incorrect."
            )
            return  # Proceed without specific remapping if config is missing

        base_style_key = "list_item_number" if ordered else "list_item_bullet"
        config_style_name = config_para_styles.get(base_style_key)

        base_style_candidate = "List Number" if ordered else "List Bullet"
        level_suffix = "" if level == 0 else f" {min(level + 1, 3)}"

        style = _best_style(
            self.styles["paragraph"],
            "paragraph",
            config_style_name,  # Prioritize config style
            f"{base_style_candidate}{level_suffix}",
            base_style_candidate,
            "Normal",
        )

        for child in element.children:
            if not hasattr(child, "name") or child.name != "li":
                continue

            # Collect nested sub-list (if any)
            nested = child.find(["ul", "ol"])

            p = self.doc.add_paragraph(style=style)
            if level > 0:
                p.paragraph_format.left_indent = Inches(0.25 * (level + 1))

            # Set list numbering properties
            p_pr = p._p.get_or_add_pPr()
            num_pr = p_pr.get_or_add_numPr()
            num_pr.get_or_add_ilvl().val = level  # Set indentation level
            num_pr.get_or_add_numId().val = (
                decimal_num_id if ordered else bullet_num_id
            )  # Set numId from config

            # Add li text, skipping nested list nodes
            for item_child in child.children:
                if hasattr(item_child, "name") and item_child.name in ("ul", "ol"):
                    continue
                if isinstance(item_child, NavigableString):
                    text = str(item_child)
                    if text.strip():
                        _add_run_with_fmt(document, available_char_styles, p, text)
                elif hasattr(item_child, "name"):
                    _inline(
                        document,
                        available_char_styles,
                        p,
                        item_child,
                        config_char_styles,
                        None,
                    )

            # Recurse into nested list
            if nested:
                self._process_list(
                    nested,
                    ordered=(nested.name == "ol"),
                    level=level + 1,
                    document=document,
                    available_char_styles=available_char_styles,
                    config_para_styles=config_para_styles,
                    config_char_styles=config_char_styles,
                )

    def _process_table(
        self,
        element,
        document: Document,
        available_char_styles: set,
        config_char_styles: Dict[str, str],
    ) -> None:
        rows = element.find_all("tr")
        if not rows:
            return

        max_cols = max((len(r.find_all(["td", "th"])) for r in rows), default=0)
        if max_cols == 0:
            return

        table = self.doc.add_table(rows=len(rows), cols=max_cols)

        # Apply a table style if one exists
        for candidate in (
            "Table Grid",
            "Light Grid",
            "Medium Grid 1",
            "Medium Shading 1 Accent 1",
            "Table Normal",
        ):
            if candidate in self.styles["paragraph"]:
                try:
                    table.style = candidate
                except Exception:
                    pass
                break

        for row_idx, tr in enumerate(rows):
            cells = tr.find_all(["td", "th"])
            for col_idx, cell_elem in enumerate(cells):
                if col_idx >= max_cols:
                    break
                cell = table.cell(row_idx, col_idx)
                # Clear the auto-created empty paragraph
                cell.paragraphs[0].clear()
                p = cell.paragraphs[0]
                _inline(
                    document,
                    available_char_styles,
                    p,
                    cell_elem,
                    config_char_styles,
                    None,
                )
                if cell_elem.name == "th":
                    for run in p.runs:
                        run.bold = True

    def _add_paragraph(self, text: str, style: str) -> None:
        resolved = _best_style(self.styles["paragraph"], "paragraph", style, "Normal")
        self.doc.add_paragraph(text, style=resolved)
