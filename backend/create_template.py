"""
Generates templates/corporate_template.docx — a styled Word template that
demonstrates what the document processor can preserve:

  - Portrait A4 page with standard margins
  - Custom heading styles  (H1–H4) in corporate blue (#1F3864)
  - "Body Text" style with slightly larger line spacing
  - "Code" style (Courier New, grey background shading)
  - "Quote" and "Intense Quote" styles
  - Header: company logo (generated PNG) + company name
  - Footer: document title centred, page number on the right
  - Table Grid style (already built-in, referenced explicitly)

Run:
    python create_template.py
"""

import io
import sys
from pathlib import Path

# ---------------------------------------------------------------------------
# Dependency check
# ---------------------------------------------------------------------------
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.oxml.ns import nsmap
except ImportError:
    sys.exit("python-docx not installed.  Run: pip install python-docx")

try:
    from PIL import Image, ImageDraw, ImageFont
    HAS_PILLOW = True
except ImportError:
    HAS_PILLOW = False
    print("Pillow not installed — header logo will be text-only.")

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
CORPORATE_BLUE = RGBColor(0x1F, 0x38, 0x64)   # dark navy
ACCENT_BLUE    = RGBColor(0x2E, 0x75, 0xB6)   # mid blue
LIGHT_GREY     = RGBColor(0xF2, 0xF2, 0xF2)
WHITE          = RGBColor(0xFF, 0xFF, 0xFF)
CODE_BG        = RGBColor(0xF4, 0xF4, 0xF4)

OUTPUT_PATH = Path(__file__).parent.parent / "templates" / "corporate_template.docx"
OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)


# ---------------------------------------------------------------------------
# Logo image builder (Pillow)
# ---------------------------------------------------------------------------

def _make_logo_png() -> io.BytesIO:
    """Return a 480×80 PNG with the company wordmark."""
    W, H = 480, 80
    img = Image.new("RGBA", (W, H), (31, 56, 100, 255))   # corporate blue
    draw = ImageDraw.Draw(img)

    # Accent bar on the left
    draw.rectangle([(0, 0), (8, H)], fill=(46, 117, 182, 255))

    # Company name
    try:
        # Try to load a system font; fall back to default
        font_big = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 28)
        font_small = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 16)
    except Exception:
        font_big = ImageFont.load_default()
        font_small = font_big

    draw.text((20, 10), "ACME CORPORATION", fill=(255, 255, 255, 255), font=font_big)
    draw.text((20, 50), "Professional Services Division", fill=(180, 210, 240, 255), font=font_small)

    buf = io.BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------

def _set_font(run_or_font, name: str, size_pt: float,
              bold: bool = False, italic: bool = False,
              color: RGBColor = None) -> None:
    font = run_or_font if hasattr(run_or_font, "name") else run_or_font.font
    font.name = name
    font.size = Pt(size_pt)
    font.bold = bold
    font.italic = italic
    if color:
        font.color.rgb = color


def _set_paragraph_spacing(fmt, before_pt: float = 0, after_pt: float = 0,
                            line_pt: float = None) -> None:
    fmt.space_before = Pt(before_pt)
    fmt.space_after = Pt(after_pt)
    if line_pt:
        from docx.shared import Pt as _Pt
        fmt.line_spacing = _Pt(line_pt)


def _shade_paragraph(paragraph, hex_color: str) -> None:
    """Apply a solid background shading to a paragraph (for Code style)."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def _add_page_border(section) -> None:
    """Add a subtle single-line page border."""
    sect_pr = section._sectPr
    pg_borders = OxmlElement("w:pgBorders")
    pg_borders.set(qn("w:offsetFrom"), "page")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "24")
        border.set(qn("w:color"), "2E75B6")
        pg_borders.append(border)
    sect_pr.append(pg_borders)


# ---------------------------------------------------------------------------
# Header builder
# ---------------------------------------------------------------------------

def _build_header(section, logo_buf: io.BytesIO = None) -> None:
    header = section.header
    header.is_linked_to_previous = False

    # Use a 1×2 table so the logo is left and the tagline is right
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))

    left_cell  = table.cell(0, 0)
    right_cell = table.cell(0, 1)

    # Left: logo image (or fallback text)
    left_p = left_cell.paragraphs[0]
    left_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if logo_buf:
        run = left_p.add_run()
        run.add_picture(logo_buf, height=Inches(0.6))
    else:
        run = left_p.add_run("ACME CORPORATION")
        _set_font(run, "Arial", 18, bold=True, color=CORPORATE_BLUE)

    # Right: subtitle text
    right_p = right_cell.paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = right_p.add_run("Professional Services")
    _set_font(run, "Arial", 9, italic=True, color=ACCENT_BLUE)
    right_p.add_run("\nDocument Management System")

    # Horizontal rule below header
    rule_p = header.add_paragraph()
    pPr = rule_p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E75B6")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ---------------------------------------------------------------------------
# Footer builder
# ---------------------------------------------------------------------------

def _build_footer(section) -> None:
    footer = section.footer
    footer.is_linked_to_previous = False

    table = footer.add_table(rows=1, cols=3, width=Inches(6.5))

    left_p   = table.cell(0, 0).paragraphs[0]
    center_p = table.cell(0, 1).paragraphs[0]
    right_p  = table.cell(0, 2).paragraphs[0]

    left_p.alignment   = WD_ALIGN_PARAGRAPH.LEFT
    center_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    right_p.alignment  = WD_ALIGN_PARAGRAPH.RIGHT

    run_l = left_p.add_run("ACME CORPORATION — Confidential")
    _set_font(run_l, "Arial", 8, color=ACCENT_BLUE)

    run_c = center_p.add_run("Document Management System")
    _set_font(run_c, "Arial", 8, color=CORPORATE_BLUE)

    # Page number field
    run_r = right_p.add_run("Page ")
    _set_font(run_r, "Arial", 8)
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    right_p._p.append(fld)
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    right_p._p.append(instr)
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    right_p._p.append(fld2)
    run_r2 = right_p.add_run(" of ")
    _set_font(run_r2, "Arial", 8)
    fld3 = OxmlElement("w:fldChar")
    fld3.set(qn("w:fldCharType"), "begin")
    right_p._p.append(fld3)
    instr2 = OxmlElement("w:instrText")
    instr2.text = "NUMPAGES"
    right_p._p.append(instr2)
    fld4 = OxmlElement("w:fldChar")
    fld4.set(qn("w:fldCharType"), "end")
    right_p._p.append(fld4)


# ---------------------------------------------------------------------------
# Style customisations
# ---------------------------------------------------------------------------

def _customise_styles(doc: Document) -> None:
    styles = doc.styles

    # ---- Heading 1 ----
    h1 = styles["Heading 1"]
    h1.font.name  = "Arial"
    h1.font.size  = Pt(22)
    h1.font.bold  = True
    h1.font.color.rgb = CORPORATE_BLUE
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after  = Pt(6)
    h1.paragraph_format.keep_with_next = True
    # Bottom border
    pPr = h1._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E75B6")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # ---- Heading 2 ----
    h2 = styles["Heading 2"]
    h2.font.name  = "Arial"
    h2.font.size  = Pt(16)
    h2.font.bold  = True
    h2.font.color.rgb = ACCENT_BLUE
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after  = Pt(4)
    h2.paragraph_format.keep_with_next = True

    # ---- Heading 3 ----
    h3 = styles["Heading 3"]
    h3.font.name   = "Arial"
    h3.font.size   = Pt(13)
    h3.font.bold   = True
    h3.font.italic = True
    h3.font.color.rgb = CORPORATE_BLUE
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after  = Pt(2)

    # ---- Heading 4 ----
    h4 = styles["Heading 4"]
    h4.font.name  = "Arial"
    h4.font.size  = Pt(11)
    h4.font.bold  = True
    h4.font.color.rgb = ACCENT_BLUE
    h4.paragraph_format.space_before = Pt(8)
    h4.paragraph_format.space_after  = Pt(2)

    # ---- Normal / Body ----
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after  = Pt(6)
    normal.paragraph_format.line_spacing = Pt(14)

    # ---- Body Text (used for regular paragraphs) ----
    try:
        body_text = styles["Body Text"]
    except KeyError:
        body_text = styles.add_style("Body Text", 1)  # 1 = paragraph
    body_text.base_style = styles["Normal"]
    body_text.font.name = "Arial"
    body_text.font.size = Pt(11)
    body_text.paragraph_format.space_after  = Pt(8)
    body_text.paragraph_format.line_spacing = Pt(15)

    # ---- Quote ----
    try:
        quote = styles["Quote"]
    except KeyError:
        quote = styles.add_style("Quote", 1)
    quote.font.name   = "Georgia"
    quote.font.size   = Pt(11)
    quote.font.italic = True
    quote.font.color.rgb = ACCENT_BLUE
    quote.paragraph_format.left_indent  = Inches(0.5)
    quote.paragraph_format.right_indent = Inches(0.5)
    quote.paragraph_format.space_before = Pt(6)
    quote.paragraph_format.space_after  = Pt(6)

    # ---- Intense Quote ----
    try:
        iq = styles["Intense Quote"]
    except KeyError:
        iq = styles.add_style("Intense Quote", 1)
    iq.base_style = styles["Quote"]
    iq.font.bold  = True
    iq.font.color.rgb = CORPORATE_BLUE

    # ---- Code (monospace, grey background) ----
    try:
        code = styles["Code"]
    except KeyError:
        code = styles.add_style("Code", 1)
    code.font.name = "Courier New"
    code.font.size = Pt(9)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after  = Pt(4)
    code.paragraph_format.left_indent  = Inches(0.25)

    # ---- List Bullet ----
    for style_name in ("List Bullet", "List Bullet 2", "List Bullet 3"):
        try:
            s = styles[style_name]
            s.font.name = "Arial"
            s.font.size = Pt(11)
        except KeyError:
            pass

    # ---- List Number ----
    for style_name in ("List Number", "List Number 2", "List Number 3"):
        try:
            s = styles[style_name]
            s.font.name = "Arial"
            s.font.size = Pt(11)
        except KeyError:
            pass

    # ---- Caption ----
    try:
        caption = styles["Caption"]
        caption.font.name   = "Arial"
        caption.font.size   = Pt(9)
        caption.font.italic = True
        caption.font.color.rgb = ACCENT_BLUE
    except KeyError:
        pass


# ---------------------------------------------------------------------------
# Page setup
# ---------------------------------------------------------------------------

def _setup_section(section) -> None:
    section.orientation   = WD_ORIENT.PORTRAIT
    section.page_width    = Cm(21)     # A4
    section.page_height   = Cm(29.7)
    section.left_margin   = Cm(2.54)
    section.right_margin  = Cm(2.54)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.header_distance = Cm(1.27)
    section.footer_distance = Cm(1.27)


# ---------------------------------------------------------------------------
# Cover / sample content
# ---------------------------------------------------------------------------

def _add_cover_content(doc: Document) -> None:
    """
    Add a minimal placeholder body so the template is not entirely empty.
    The document processor will replace this with real content.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[Document Title]")
    run.font.name  = "Arial"
    run.font.size  = Pt(28)
    run.font.bold  = True
    run.font.color.rgb = CORPORATE_BLUE

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub.add_run("[Subtitle / Document Reference]")
    run2.font.name   = "Arial"
    run2.font.size   = Pt(14)
    run2.font.color.rgb = ACCENT_BLUE

    doc.add_paragraph()  # spacer

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = meta.add_run("Prepared by: ACME Corporation\nDate: [Date]")
    run3.font.name  = "Arial"
    run3.font.size  = Pt(11)
    run3.font.color.rgb = CORPORATE_BLUE

    doc.add_page_break()

    toc_heading = doc.add_paragraph("Table of Contents", style="Heading 1")
    doc.add_paragraph(
        "This section is a placeholder. The document processor will populate "
        "this template with the supplied Markdown content.",
        style="Body Text",
    )


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def create_template(output_path: Path = OUTPUT_PATH) -> Path:
    doc = Document()

    # Page layout
    section = doc.sections[0]
    _setup_section(section)
    _add_page_border(section)

    # Styles
    _customise_styles(doc)

    # Header / footer
    logo_buf = _make_logo_png() if HAS_PILLOW else None
    _build_header(section, logo_buf)
    _build_footer(section)

    # Placeholder body
    _add_cover_content(doc)

    doc.save(output_path)
    print(f"Template written to: {output_path}")
    return output_path


if __name__ == "__main__":
    create_template()
